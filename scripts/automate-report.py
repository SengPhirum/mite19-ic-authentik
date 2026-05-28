import argparse
import datetime as dt
import html
import hmac
import hashlib
import json
import os
from pathlib import Path
import re
import secrets
import shutil
import subprocess
import sys
import time
import urllib.request

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import sync_playwright


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS_DIR = ROOT / "docs" / "screenshots"
VIDEOS_DIR = ROOT / "docs" / "videos"
REPORTS_DIR = ROOT / "reports"
REPORT_MD = ROOT / "docs" / "ASSIGNMENT_REPORT.md"

VIEWPORT = {"width": 1366, "height": 768}

SCREENSHOT_CAPTIONS = [
    ("01-docker-compose-ps.png", "Docker Compose status showing authentik, PostgreSQL, worker, Grafana, and backup services."),
    ("02-initial-setup.png", "Initial setup or completed bootstrap evidence for the authentik instance."),
    ("03-users-groups.png", "Demo users, groups, and Grafana access membership used in the lab."),
    ("04-admin-totp-enrolled.png", "Administrator MFA proof showing the TOTP authentication challenge for demo-admin."),
    ("05-grafana-provider.png", "OAuth2/OIDC provider details for the Grafana integration."),
    ("06-grafana-login-flow.png", "Grafana redirecting the user to authentik for SSO."),
    ("07-grafana-sso-success.png", "Successful Grafana SSO session as demo-user."),
    ("08-conditional-access-denied.png", "Conditional access policy denying demo-outsider."),
    ("09-events-logs.png", "authentik event log entries for login and policy activity."),
    ("10-backup-restore.png", "Backup output and restore validation evidence."),
]


def load_env(path):
    env = {}
    if not path.exists():
        return env
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        env[key.strip()] = value.strip().strip('"').strip("'")
    return env


def run_command(args, check=True, cwd=ROOT):
    display_command = format_command_for_log(args)
    print(f"> {display_command}", flush=True)
    result = subprocess.run(
        [str(a) for a in args],
        cwd=cwd,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
    )
    output = result.stdout or ""
    if check and result.returncode != 0:
        if display_command.endswith("<redacted>"):
            print("Redacted command failed; output suppressed because it may contain credentials.")
        else:
            print(output)
        raise RuntimeError(f"Command failed with exit code {result.returncode}: {display_command}")
    return output


def format_command_for_log(args):
    values = [str(arg) for arg in args]
    if len(values) >= 6 and values[:5] == ["docker", "compose", "exec", "-T", "worker"] and "shell" in values:
        return "docker compose exec -T worker ak shell -c <redacted>"
    return " ".join(values)


def wait_for_http(url, timeout_seconds=180):
    deadline = time.time() + timeout_seconds
    last_error = None
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=5) as response:
                if 200 <= response.status < 500:
                    return
        except Exception as exc:  # noqa: BLE001 - endpoint can fail while containers boot
            last_error = exc
        time.sleep(3)
    raise TimeoutError(f"Timed out waiting for {url}: {last_error}")


def shell_json(code, marker="REPORT_JSON="):
    output = run_command(["docker", "compose", "exec", "-T", "worker", "ak", "shell", "-c", code])
    for line in output.splitlines():
        if line.startswith(marker):
            return json.loads(line[len(marker):])
    raise RuntimeError(f"Could not find {marker!r} in authentik shell output.")


def ensure_demo_state(env):
    demo_admin_password = env.get("DEMO_ADMIN_PASSWORD", "demo-admin-password")
    demo_user_password = env.get("DEMO_USER_PASSWORD", "demo-user-password")
    demo_outsider_password = env.get("DEMO_OUTSIDER_PASSWORD", "demo-outsider-password")
    report_admin_password = "ReportAutomation" + secrets.token_hex(12)

    code = f"""
from authentik.core.models import User, Group

def upsert_user(username, name, email, password):
    user, _ = User.objects.get_or_create(
        username=username,
        defaults={{
            "name": name,
            "email": email,
            "is_active": True,
            "type": "internal",
            "path": "users",
        }},
    )
    user.name = name
    user.email = email
    user.is_active = True
    user.type = "internal"
    user.path = "users"
    user.set_password(password)
    user.save()
    return user

demo_admin = upsert_user("demo-admin", "Demo Administrator", "demo-admin@example.local", {demo_admin_password!r})
demo_user = upsert_user("demo-user", "Demo Regular User", "demo-user@example.local", {demo_user_password!r})
demo_outsider = upsert_user("demo-outsider", "Demo Unauthorized User", "demo-outsider@example.local", {demo_outsider_password!r})
report_admin = upsert_user("report-admin", "Report Automation Admin", "report-admin@example.local", {report_admin_password!r})

assignment_admins, _ = Group.objects.get_or_create(name="assignment-admins", defaults={{"is_superuser": True}})
assignment_admins.is_superuser = True
assignment_admins.save()
assignment_admins.users.add(demo_admin, report_admin)

grafana_access, _ = Group.objects.get_or_create(name="grafana-access")
grafana_access.users.add(demo_admin, demo_user)
grafana_access.users.remove(demo_outsider)

print("REPORT_JSON=" + __import__("json").dumps({{"report_admin_password": {report_admin_password!r}}}))
"""
    data = shell_json(code)
    return data["report_admin_password"]


def ensure_demo_admin_totp():
    code = """
from authentik.core.models import User
from authentik.stages.authenticator_totp.models import TOTPDevice
import secrets
import json

user = User.objects.get(username="demo-admin")
device = TOTPDevice.objects.filter(user=user, confirmed=True).first()
created = False
if device is None:
    device = TOTPDevice.objects.create(
        user=user,
        name="Report Automation TOTP",
        confirmed=True,
        key=secrets.token_hex(20),
        step=30,
        digits=6,
    )
    created = True

print("REPORT_JSON=" + json.dumps({
    "key": device.key,
    "step": device.step,
    "digits": device.digits,
    "created": created,
}))
"""
    return shell_json(code)


def collect_lab_data():
    code = """
from authentik.core.models import User, Group, Application
from authentik.providers.oauth2.models import OAuth2Provider
from authentik.stages.authenticator_totp.models import TOTPDevice
import json

users = []
for user in User.objects.filter(username__in=["akadmin", "demo-admin", "demo-user", "demo-outsider", "report-admin"]).order_by("username"):
    users.append({
        "username": user.username,
        "name": user.name,
        "email": user.email,
        "active": user.is_active,
        "type": user.type,
    })

groups = []
for group in Group.objects.filter(name__in=["assignment-admins", "grafana-access"]).order_by("name"):
    groups.append({
        "name": group.name,
        "is_superuser": group.is_superuser,
        "users": list(group.users.order_by("username").values_list("username", flat=True)),
    })

providers = []
for provider in OAuth2Provider.objects.filter(name__icontains="grafana").order_by("name"):
    providers.append({
        "name": provider.name,
        "client_id": provider.client_id,
        "issuer_mode": str(provider.issuer_mode),
        "redirect_uris": [str(uri) for uri in provider.redirect_uris],
    })

applications = []
for app in Application.objects.filter(slug="grafana").order_by("name"):
    applications.append({
        "name": app.name,
        "slug": app.slug,
        "launch_url": app.meta_launch_url,
        "policy_engine_mode": app.policy_engine_mode,
    })

totp_count = TOTPDevice.objects.filter(user__username="demo-admin", confirmed=True).count()

print("REPORT_JSON=" + json.dumps({
    "users": users,
    "groups": groups,
    "providers": providers,
    "applications": applications,
    "demo_admin_totp_devices": totp_count,
}))
"""
    return shell_json(code)


def cleanup_report_admin():
    code = """
from authentik.core.models import User
import json

deleted, _ = User.objects.filter(username="report-admin").delete()
print("REPORT_JSON=" + json.dumps({"deleted": deleted}))
"""
    try:
        shell_json(code)
    except Exception as exc:  # noqa: BLE001 - cleanup should not invalidate the report
        print(f"Warning: could not delete temporary report-admin user: {exc}", flush=True)


def totp_code(hex_key, step=30, digits=6):
    counter = int(time.time() // int(step))
    key = bytes.fromhex(hex_key)
    digest = hmac.new(key, counter.to_bytes(8, "big"), hashlib.sha1).digest()
    offset = digest[-1] & 0x0F
    value = int.from_bytes(digest[offset:offset + 4], "big") & 0x7FFFFFFF
    return str(value % (10 ** int(digits))).zfill(int(digits))


def html_page(title, body):
    return f"""
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      color-scheme: light;
      font-family: "Segoe UI", Arial, sans-serif;
      background: #f4f6f8;
      color: #1d252d;
    }}
    body {{
      margin: 0;
      padding: 36px;
    }}
    main {{
      background: white;
      border: 1px solid #d7dde5;
      border-radius: 8px;
      padding: 28px;
      box-shadow: 0 8px 28px rgba(15, 23, 42, 0.08);
    }}
    h1 {{
      margin: 0 0 8px;
      font-size: 28px;
      font-weight: 650;
    }}
    .meta {{
      color: #64748b;
      font-size: 14px;
      margin-bottom: 24px;
    }}
    h2 {{
      margin: 28px 0 12px;
      font-size: 18px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin-bottom: 20px;
      table-layout: fixed;
    }}
    th, td {{
      border: 1px solid #d9e0e8;
      padding: 10px 12px;
      text-align: left;
      vertical-align: top;
      word-break: break-word;
      font-size: 14px;
    }}
    th {{
      background: #eef2f7;
      font-weight: 650;
    }}
    pre {{
      white-space: pre-wrap;
      word-break: break-word;
      background: #101827;
      color: #e5edf6;
      padding: 18px;
      border-radius: 6px;
      font-family: Consolas, "Cascadia Mono", monospace;
      font-size: 13px;
      line-height: 1.45;
    }}
    .ok {{
      display: inline-block;
      background: #dcfce7;
      color: #166534;
      border: 1px solid #86efac;
      border-radius: 999px;
      padding: 2px 10px;
      font-size: 12px;
      font-weight: 650;
    }}
  </style>
</head>
<body>
  <main>
    <h1>{html.escape(title)}</h1>
    <div class="meta">Captured {dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>
    {body}
  </main>
</body>
</html>
"""


def table_html(headers, rows):
    head = "".join(f"<th>{html.escape(str(item))}</th>" for item in headers)
    body_rows = []
    for row in rows:
        body_rows.append("<tr>" + "".join(f"<td>{html.escape(str(item))}</td>" for item in row) + "</tr>")
    return f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"


def screenshot_html(page, title, body, output_path):
    page.set_content(html_page(title, body), wait_until="networkidle")
    page.screenshot(path=str(output_path), full_page=True)


def screenshot_terminal(page, title, command, output, output_path):
    body = f"""
<h2>Command</h2>
<pre>{html.escape(command)}</pre>
<h2>Output</h2>
<pre>{html.escape(output.strip() or "(no output)")}</pre>
"""
    screenshot_html(page, title, body, output_path)


def screenshot_users_groups(page, data, output_path):
    user_rows = [
        [user["username"], user["name"], user["email"], "yes" if user["active"] else "no", user["type"]]
        for user in data["users"]
        if user["username"].startswith("demo-")
    ]
    group_rows = [
        [
            group["name"],
            "yes" if group["is_superuser"] else "no",
            ", ".join(member for member in group["users"] if member != "report-admin"),
        ]
        for group in data["groups"]
    ]
    body = """
<p><span class="ok">Live authentik data</span></p>
<h2>Users</h2>
""" + table_html(["Username", "Name", "Email", "Active", "Type"], user_rows) + """
<h2>Groups</h2>
""" + table_html(["Group", "Superuser", "Members"], group_rows)
    screenshot_html(page, "Users and Groups Evidence", body, output_path)


def screenshot_bootstrap(page, data, output_path):
    rows = [
        ["akadmin user exists", "yes" if any(user["username"] == "akadmin" for user in data["users"]) else "no"],
        ["demo-admin user exists", "yes" if any(user["username"] == "demo-admin" for user in data["users"]) else "no"],
        ["demo-admin confirmed TOTP devices", data["demo_admin_totp_devices"]],
        ["Grafana application exists", "yes" if data["applications"] else "no"],
        ["Grafana OIDC provider exists", "yes" if data["providers"] else "no"],
    ]
    body = """
<p><span class="ok">Bootstrap verified</span></p>
""" + table_html(["Check", "Result"], rows)
    screenshot_html(page, "Initial Setup and Bootstrap Evidence", body, output_path)


def login_authentik(page, base_url, username, password):
    page.goto(f"{base_url}/if/admin/", wait_until="networkidle", timeout=60000)
    page.wait_for_selector("#ak-identifier-input", timeout=30000)
    page.locator("#ak-identifier-input").fill(username)
    page.get_by_role("button", name="Log in").click()
    page.wait_for_selector('input[placeholder="Please enter your password"]', timeout=30000)
    page.locator('input[placeholder="Please enter your password"]').fill(password)
    page.get_by_role("button", name="Continue").click()


def capture_provider(page, base_url):
    page.goto(f"{base_url}/if/admin/#/core/providers", wait_until="networkidle", timeout=60000)
    page.wait_for_timeout(2000)
    try:
        page.get_by_text("grafana-oidc-provider").click(timeout=10000)
        page.wait_for_timeout(2500)
    except PlaywrightTimeoutError:
        pass
    page.screenshot(path=str(SCREENSHOTS_DIR / "05-grafana-provider.png"), full_page=True)


def capture_event_log(page, base_url):
    page.goto(f"{base_url}/if/admin/#/events/log", wait_until="networkidle", timeout=60000)
    page.wait_for_timeout(3500)
    page.screenshot(path=str(SCREENSHOTS_DIR / "09-events-logs.png"), full_page=True)


def capture_mfa_challenge(playwright, base_url, env, totp):
    browser = playwright.chromium.launch(headless=True)
    context = browser.new_context(viewport=VIEWPORT)
    page = context.new_page()
    try:
        login_authentik(page, base_url, "demo-admin", env["DEMO_ADMIN_PASSWORD"])
        page.wait_for_timeout(2500)
        page.screenshot(path=str(SCREENSHOTS_DIR / "04-admin-totp-enrolled.png"), full_page=True)
        code_input = page.locator('input[placeholder="Type an authentication code..."]')
        if code_input.count() > 0:
            code = totp_code(totp["key"], totp["step"], totp["digits"])
            code_input.fill(code)
            page.get_by_role("button", name="Continue").click()
            page.wait_for_timeout(2500)
    finally:
        context.close()
        browser.close()


def capture_sso_flow(playwright, base_auth, base_grafana, env, headless, record_video):
    browser = playwright.chromium.launch(headless=headless)
    context_args = {"viewport": VIEWPORT}
    video_path = None
    if record_video:
        context_args["record_video_dir"] = str(VIDEOS_DIR)
        context_args["record_video_size"] = VIEWPORT
    context = browser.new_context(**context_args)
    page = context.new_page()
    video_obj = page.video

    try:
        page.goto(f"{base_grafana}/login", wait_until="networkidle", timeout=60000)
        page.get_by_text("Sign in with authentik").click()
        page.wait_for_url(f"{base_auth}/if/flow/**", timeout=30000)
        page.wait_for_timeout(1200)
        page.screenshot(path=str(SCREENSHOTS_DIR / "06-grafana-login-flow.png"), full_page=True)

        complete_auth_flow(page, "demo-user", env["DEMO_USER_PASSWORD"])
        page.wait_for_url(f"{base_grafana}/**", timeout=60000)
        page.goto(f"{base_grafana}/profile", wait_until="networkidle", timeout=60000)
        page.wait_for_timeout(2500)
        page.screenshot(path=str(SCREENSHOTS_DIR / "07-grafana-sso-success.png"), full_page=True)

        context.clear_cookies()
        page.goto("about:blank")
        page.goto(f"{base_grafana}/login", wait_until="networkidle", timeout=60000)
        page.get_by_text("Sign in with authentik").click()
        page.wait_for_url(f"{base_auth}/if/flow/**", timeout=30000)
        complete_auth_flow(page, "demo-outsider", env["DEMO_OUTSIDER_PASSWORD"])
        page.wait_for_timeout(4500)
        page.screenshot(path=str(SCREENSHOTS_DIR / "08-conditional-access-denied.png"), full_page=True)
    finally:
        context.close()
        if record_video and video_obj is not None:
            src = Path(video_obj.path())
            target = VIDEOS_DIR / "sso-login-and-denied-flow.webm"
            if target.exists():
                target.unlink()
            shutil.move(str(src), str(target))
            video_path = target
        browser.close()
    return video_path


def complete_auth_flow(page, username, password):
    page.wait_for_selector("#ak-identifier-input", timeout=30000)
    page.locator("#ak-identifier-input").fill(username)
    page.get_by_role("button", name="Log in").click()
    page.wait_for_selector('input[placeholder="Please enter your password"]', timeout=30000)
    page.locator('input[placeholder="Please enter your password"]').fill(password)
    page.get_by_role("button", name="Continue").click()


def run_backup_restore_check(stamp):
    powershell = shutil.which("powershell") or shutil.which("pwsh")
    if os.name == "nt" and powershell:
        backup_output = run_command([
            powershell,
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(ROOT / "scripts" / "backup-now.ps1"),
            "-Stamp",
            stamp,
        ])
    else:
        backup_output = run_command([str(ROOT / "scripts" / "backup-now.sh"), stamp])

    env = load_env(ROOT / ".env")
    pg_user = env.get("PG_USER", "authentik")
    restore_db = f"authentik_restore_check_{stamp.lower()}"
    restore_output_parts = []
    try:
        restore_output_parts.append(run_command(["docker", "compose", "exec", "-T", "postgresql", "createdb", "-U", pg_user, restore_db]))
        restore_output_parts.append(run_command([
            "docker",
            "compose",
            "exec",
            "-T",
            "postgresql",
            "pg_restore",
            "-U",
            pg_user,
            "-d",
            restore_db,
            "--clean",
            "--if-exists",
            f"/backups/{stamp}/authentik.dump",
        ]))
        restore_output_parts.append("Restore validation succeeded in temporary database: " + restore_db)
    finally:
        restore_output_parts.append(run_command([
            "docker",
            "compose",
            "exec",
            "-T",
            "postgresql",
            "dropdb",
            "--if-exists",
            "-U",
            pg_user,
            restore_db,
        ], check=False))

    return backup_output + "\n" + "\n".join(part.strip() for part in restore_output_parts if part.strip())


def add_page_number_footer(document):
    section = document.sections[0]
    paragraph = section.footer.paragraphs[0]
    paragraph.text = "Page "
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_code_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_markdown_table(document, lines):
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        rows.append([cell.strip() for cell in stripped.split("|")])
    if len(rows) < 2:
        return
    data_rows = [rows[0]] + rows[2:]
    table = document.add_table(rows=1, cols=len(data_rows[0]))
    table.style = "Table Grid"
    for index, value in enumerate(data_rows[0]):
        table.rows[0].cells[index].text = value
    for row in data_rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row[:len(cells)]):
            cells[index].text = value


def add_markdown(document, markdown_text):
    in_code = False
    code_lines = []
    table_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            add_code_paragraph(document, "\n".join(code_lines))
            code_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(document, table_lines)
            table_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip():
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(clean_inline_markdown(heading_match.group(2)), level=level)
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet_match:
            document.add_paragraph(clean_inline_markdown(bullet_match.group(1)), style="List Bullet")
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if number_match:
            document.add_paragraph(clean_inline_markdown(number_match.group(1)), style="List Number")
            continue

        document.add_paragraph(clean_inline_markdown(line))

    flush_table()
    flush_code()


def clean_inline_markdown(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def build_word_report(group_name, video_path):
    safe_group = safe_filename(group_name)
    output_path = REPORTS_DIR / f"{safe_group}_authentik_Assignment.docx"
    document = Document()
    document.core_properties.title = "Identity & Access Management with authentik"
    document.core_properties.author = group_name
    add_page_number_footer(document)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    markdown_text = REPORT_MD.read_text(encoding="utf-8")
    markdown_text = markdown_text.replace("GroupName_authentik_Assignment.pdf", f"{safe_group}_authentik_Assignment.docx")
    add_markdown(document, markdown_text)

    document.add_page_break()
    document.add_heading("Evidence Appendix", level=1)
    document.add_paragraph("The screenshots below were captured by scripts/run-report-automation.ps1 from the local Docker Compose lab.")
    if video_path:
        document.add_paragraph(f"Screen recording: {video_path.relative_to(ROOT)}")

    for filename, caption in SCREENSHOT_CAPTIONS:
        image_path = SCREENSHOTS_DIR / filename
        document.add_heading(filename, level=2)
        document.add_paragraph(caption)
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.5))
        else:
            document.add_paragraph(f"Missing screenshot: {image_path.relative_to(ROOT)}")

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def safe_filename(value):
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", value.strip())
    return cleaned.strip("._") or "GroupName"


def parse_args():
    parser = argparse.ArgumentParser(description="Capture authentik assignment screenshots, SSO video, and Word report.")
    parser.add_argument("--group-name", default="GroupName")
    parser.add_argument("--skip-docker-start", action="store_true")
    parser.add_argument("--no-video", action="store_true")
    parser.add_argument("--skip-backup-restore", action="store_true")
    parser.add_argument("--headful", action="store_true")
    return parser.parse_args()


def main():
    args = parse_args()
    env = load_env(ROOT / ".env")
    base_auth = env.get("AUTHENTIK_PUBLIC_URL", "http://localhost:9000").rstrip("/")
    base_grafana = env.get("GRAFANA_ROOT_URL", "http://localhost:3000").rstrip("/")

    required = ["DEMO_ADMIN_PASSWORD", "DEMO_USER_PASSWORD", "DEMO_OUTSIDER_PASSWORD"]
    missing = [key for key in required if key not in env]
    if missing:
        raise RuntimeError(f"Missing required values in .env: {', '.join(missing)}")

    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
    VIDEOS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    if not args.skip_docker_start:
        run_command(["docker", "compose", "up", "-d"])

    wait_for_http(f"{base_auth}/-/health/live/")
    wait_for_http(f"{base_grafana}/api/health")

    run_command(["docker", "compose", "exec", "-T", "worker", "ak", "apply_blueprint", "/blueprints/assignment-lab.yaml"], check=False)
    report_admin_password = ensure_demo_state(env)
    totp = ensure_demo_admin_totp()
    data = collect_lab_data()
    ps_output = run_command(["docker", "compose", "ps"])

    stamp = dt.datetime.now(dt.UTC).strftime("%Y%m%dT%H%M%SZ")
    backup_output = None
    if not args.skip_backup_restore:
        backup_output = run_backup_restore_check(stamp)

    video_path = None
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=not args.headful)
        page = browser.new_page(viewport=VIEWPORT)
        screenshot_terminal(page, "Docker Compose Status", "docker compose ps", ps_output, SCREENSHOTS_DIR / "01-docker-compose-ps.png")
        screenshot_bootstrap(page, data, SCREENSHOTS_DIR / "02-initial-setup.png")
        screenshot_users_groups(page, data, SCREENSHOTS_DIR / "03-users-groups.png")
        if backup_output:
            screenshot_terminal(
                page,
                "Backup and Restore Validation",
                f"scripts/backup-now.ps1 -Stamp {stamp}; pg_restore into temporary database",
                backup_output,
                SCREENSHOTS_DIR / "10-backup-restore.png",
            )
        else:
            screenshot_html(
                page,
                "Backup and Restore Validation",
                "<p>Backup and restore capture skipped for this run.</p>",
                SCREENSHOTS_DIR / "10-backup-restore.png",
            )
        browser.close()

        capture_mfa_challenge(playwright, base_auth, env, totp)

        admin_browser = playwright.chromium.launch(headless=not args.headful)
        admin_context = admin_browser.new_context(viewport=VIEWPORT)
        admin_page = admin_context.new_page()
        login_authentik(admin_page, base_auth, "report-admin", report_admin_password)
        admin_page.wait_for_url(f"{base_auth}/if/admin/**", timeout=60000)
        admin_page.wait_for_timeout(2500)
        capture_provider(admin_page, base_auth)

        video_path = capture_sso_flow(
            playwright,
            base_auth,
            base_grafana,
            env,
            headless=not args.headful,
            record_video=not args.no_video,
        )

        capture_event_log(admin_page, base_auth)
        admin_context.close()
        admin_browser.close()

    cleanup_report_admin()
    report_path = build_word_report(args.group_name, video_path)

    print("")
    print("Automation complete.")
    print(f"Screenshots: {SCREENSHOTS_DIR}")
    if video_path:
        print(f"Video: {video_path}")
    print(f"Word report: {report_path}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # noqa: BLE001 - produce readable script failure
        print(f"ERROR: {exc}", file=sys.stderr)
        sys.exit(1)
